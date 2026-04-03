#!/usr/bin/env python3
"""
Build WiDS AI Forum scientific poster as .docx (python-docx, OOXML-safe tblBorders).

WiDS Apprenticeship — Final Deliverables (poster):
  - 48\" W × 36\" H landscape, 0.5\" margins
  - .docx only; 800–1,000 words body (excl. references); 4–10 figures/tables; ≥300 DPI images
  - Font guide: Title 60–120 pt; Authors 40–60 pt; Headings 30–40 pt; Body 24–30 pt;
    Captions 20–26 pt; References 16–20 pt
  - Reading order: left → right, top → bottom (see WiDS layout reference PDF)

Content source of truth: docs/WIDS_POSTER_CONTENT_FOR_WORD.md

Usage:
  pip install python-docx
  python3 scripts/build_wids_poster_docx.py
  # Output: docs/WiDS_Poster_MinutesMatter.docx
"""

from __future__ import annotations

import io
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --- Poster copy (aligned with docs/WIDS_POSTER_CONTENT_FOR_WORD.md) ---

TITLE = (
    "Minutes Matter: An Agentic, Equity-Aware Wildfire Evacuation Platform "
    "for Faster Protective Action"
)

AUTHORS = (
    "Lena Ayesh, Katie Leedom, Anisha Nannapaneni, Nadia Narayanan"
)

AFFILIATIONS = (
    "University of North Carolina at Charlotte · WiDS Apprenticeship Pod · "
    "Accenture Song (industry partner)"
)

LOGOS_NOTE = (
    "[ Logos: WiDS · UNC Charlotte · Accenture Song — per program instructions ]"
)

SUBTITLE = (
    "Minutes Matter (WildfireAlert) · https://github.com/layesh1/wildfire-app · "
    "iOS: https://github.com/anishan3213-design/minutes-matter-ios"
)

TRACK_LINE = (
    "WiDS Datathon 2026 · Track 1 — Accelerating Equitable Evacuations"
)

ABSTRACT = (
    "Wildfire evacuations fail when alerts arrive late, messages lack actionable context, "
    "and vulnerable households are underrepresented in opt-in systems. WiDS Datathon 2026 "
    "Track 1 — Accelerating Equitable Evacuations asks how to shorten the path from first "
    "credible signal to protective action while improving clarity and completion among "
    "at-risk communities. We present Minutes Matter, an operational web and mobile "
    "ecosystem that unifies WatchDuty-scale incident analytics with live operational data. "
    "The platform combines a Next.js application on Vercel, a Supabase backend with "
    "row-level security, and Flameo, a three-phase agentic AI stack (structured context "
    "without an LLM, proactive briefing, grounded chat with guardrails) powered by "
    "Anthropic Claude. Live feeds include NIFC and NASA FIRMS, FEMA National Shelter System "
    "shelters, Open-Meteo weather, FEMA National Risk Index–style equity layers, and curated "
    "hazard facilities (nuclear, chemical, LNG) for routing awareness; Google geocoding and "
    "routes support shelter ranking with fire- and hazard-avoidance flags. We implement a "
    "two-status evacuation model (home vs personal safety) for clearer household rollups for "
    "responders, Life360-style family linking, station-based responder workflows with Flameo "
    "COMMAND, and a native SwiftUI iOS client. A Kaggle companion notebook explores "
    "WatchDuty-style data, feature engineering, proximity to hazards, shelter coverage, and "
    "equity metrics. We address datathon pitfalls—inconsistent reporting, geographic bias, "
    "and weak ground truth—by stating assumptions explicitly and grounding user-facing "
    "outputs in reproducible pipelines. Significance: a tangible path from research on "
    "alert delay and silent incidents to a consent-aware, multi-role system agencies could "
    "pilot for earlier, clearer, and more equitable evacuation support."
)

INTRODUCTION = (
    "Wildfires are increasing in frequency and intensity; public demand for trusted, "
    "hyperlocal intelligence spikes during major events. In a 2021–2025 WatchDuty-style "
    "WiDS analytic frame (62,696 U.S. fire incidents), 73.5% had no public alert (46,053 "
    "“silent” events); among silent fires, only one received an evacuation order, and "
    "median time to order when orders exist was 1.1 hours (90th percentile 32.1 hours). "
    "WatchDuty and related data also show geographic and volunteer-coverage bias, so naive "
    "risk maps can mis-rank counties. Track 1 challenges teams to reduce alert lag, improve "
    "message clarity, and raise evacuation completion for vulnerable groups. We address this "
    "gap with Minutes Matter: a deployable prototype that fuses these statistics with "
    "production-style APIs, maps, and AI grounded in live context. Our design emphasizes "
    "equitable evacuation: surfacing signal gaps, clarifying next actions, and pairing "
    "household status with responder visibility under explicit consent. Figure 1 anchors "
    "national scale (incident burden, silent fraction, delay order-of-magnitude). Figure 2 "
    "summarizes user roles (evacuee, caregiver, responder) and the signal → context → "
    "action loop the product enforces."
)

DATA_METHODS = (
    "Data: WatchDuty / WiDS competition tabular data (primary for notebook exploration); "
    "live NIFC EGP and NASA FIRMS VIIRS via app proxies; FEMA NSS open shelters; Open-Meteo "
    "for fire-weather heuristics; FEMA NRI county endpoints for wildfire risk and social "
    "vulnerability themes; static hazard facilities (30 sites: nuclear, chemical, LNG); "
    "optional Google Places, Geocoding, and Routes. Methods: Python (pandas, scikit-learn, "
    "plotly) in notebooks/kaggle_wids_2026_minutes_matter.ipynb with optional "
    "minutes_matter_kaggle.py for /kaggle/input discovery; haversine proximity, feature "
    "engineering (fire weather, population exposure), and RandomForest baselines where labels "
    "exist. Application layer: TypeScript API routes; FlameoContext via GET "
    "/api/flameo/context; briefing and chat endpoints; rate limits and CSP. Figure 3 shows "
    "the data and AI pipeline. Figure 4 shows Flameo phases (A: structured JSON context; "
    "B: briefing; C: grounded chat + COMMAND)."
)

RESULTS = (
    "Empirical alert chain (WatchDuty-style WiDS frame, 62,696 fires): 41,906 fires "
    "(99.74%) with early external signal still had no evacuation action; among 298 "
    "extreme-spread fires, 211 (70.8%) had no evacuation action. Fires in high-SVI counties "
    "grew ~17% faster than average. County SVI vs. orders: high-SVI counties (SVI > 0.70) "
    "showed ~0.7% evacuation-order rate vs ~2.4% in lower-SVI counties (SVI < 0.55)—"
    "order-of-magnitude 3.4× gap—consistent with structural exclusion, not only delay. "
    "Product: the web app delivers verified-address anchoring, hazard-aware shelter cards, "
    "push escalation, and role dashboards; iOS mirrors core field flows. Notebook: "
    "geographic concentration, missingness, density-stratified signal gaps, RandomForest "
    "feature ranking (synthetic benchmark: log acres, wind speed, containment lead; re-run "
    "on competition train.csv for held-out metrics—demo labels are rule-derived). Live "
    "/api/active-fires, /api/shelters/live, /api/weather validate the deployed stack."
)

DISCUSSION = (
    "Minutes Matter maps Track 1 goals to a testable system: earlier awareness through "
    "multi-source fusion, clearer next actions via structured UI and Flameo guardrails, and "
    "equity-aware analytics that acknowledge WatchDuty coverage bias and missing ground "
    "truth. Compared to alert-only baselines, we stress actionable routing, household and "
    "family aggregation, and responder visibility under consent. Limitations: reliance on "
    "third-party API availability; sparse shelter feeds when no incident is active; "
    "generative outputs require human oversight for life-safety messaging; notebook RF "
    "metrics on synthetic / rule-labeled rows are not claimed as real-world AUC. Next steps: "
    "tighter calibration to official evacuation orders, multilingual alert templates aligned "
    "with WiDS playbook concepts, and prospective evaluation with drills or pilots. The "
    "GitHub repository and Kaggle notebook support reproducibility for judges and partners. "
    "Per WiDS instructions, this Word file will be adapted into a Figma print layout; gray "
    "figure placeholders should be replaced with numbered ≥300 DPI exports before or after handoff."
)

# Placeholder cells for ≥300 DPI raster inserts (Figma/print team)
FIGURE_PLACEHOLDERS = {
    1: "[ Figure 1 — INSERT ≥300 DPI — National burden / silent fraction / delay (~10 in wide) ]",
    2: "[ Figure 2 — INSERT ≥300 DPI — User journey & roles (~8–10 in wide) ]",
    3: "[ Figure 3 — INSERT ≥300 DPI — End-to-end pipeline (~10 in wide) ]",
    4: "[ Figure 4 — INSERT ≥300 DPI — Flameo A/B/C + COMMAND (~8 in wide) ]",
    5: "[ Figure 5 — INSERT ≥300 DPI — Minutes Matter UI screenshot (~12 in wide) ]",
    6: "[ Figure 6 — INSERT ≥300 DPI — Notebook: incidents by state or time (~9 in wide) ]",
    7: "[ Figure 7 — INSERT ≥300 DPI — RandomForest feature importance (~9 in wide) ]",
    8: "[ Figure 8 — INSERT ≥300 DPI — Equity / delay analysis (~9 in wide) ]",
}

CAPTIONS = {
    1: (
        "Figure 1. National wildfire burden and alerting gap (2021–2025 analytic frame): "
        "62,696 incidents; 73.5% silent (46,053); median / 90th-percentile hours to order "
        "(1.1 h; 32.1 h); peak month and hour where shown."
    ),
    2: (
        "Figure 2. User journey and roles: consumers, caregivers, responders, analysts; "
        "signal → context → action; Flameo COMMAND for responders."
    ),
    3: (
        "Figure 3. End-to-end pipeline: NIFC / FIRMS / NSS / Open-Meteo / NRI / hazards / "
        "(optional) Google → Next.js APIs → Supabase + maps + Flameo."
    ),
    4: (
        "Figure 4. Flameo architecture: Phase A structured JSON (no LLM); B briefing; "
        "C grounded chat + guardrails; COMMAND mode."
    ),
    5: (
        "Figure 5. Minutes Matter UI: consumer hub or evacuation map with shelters, fire "
        "layers, hazard context."
    ),
    6: "Figure 6. Notebook: WatchDuty-style exploration — incidents by state or time.",
    7: (
        "Figure 7. Notebook: RandomForest feature importance (log acres, wind, containment "
        "typical on demo; update after competition train run)."
    ),
    8: (
        "Figure 8. Notebook: equity / delay — e.g. silent rate by density quintile or "
        "SVI vs wildfire risk by county."
    ),
}

TABLE1_ROWS = [
    ("#", "Source", "Role"),
    ("1", "WatchDuty / WiDS", "Primary tabular incidents and labels (notebook)"),
    ("2", "NASA FIRMS (VIIRS)", "Live hotspots (/api/fires/firms)"),
    ("3", "NIFC (EGP / WFIGS)", "Active incidents (/api/active-fires, /api/fires/nifc)"),
    ("4", "FEMA NSS", "Open shelters (/api/shelters/live)"),
    ("5", "FEMA NRI", "County wildfire risk & SVI (/api/nri)"),
    ("6", "Open-Meteo", "Weather and fire-weather heuristic (/api/weather)"),
    ("7", "Hazard facilities (30)", "Nuclear / chemical / LNG — routing and context"),
    ("8", "Google Places / Routes", "Geocoding and hazard-aware shelter routing"),
]

TABLE2_ROWS = [
    ("Theme", "Result / demonstration"),
    ("Silent / signal gap", "73.5% no public alert; 99.74% no action despite early signal"),
    ("Equity", "3.4× order-rate gap high- vs lower-SVI; +17% faster growth in vulnerable counties"),
    ("Extreme events", "70.8% of extreme-spread fires with no evacuation action"),
    ("Flameo / routing", "Phase A JSON context; shelter routes with hazard-buffer awareness"),
    ("Responder rollup", "Two-status model (home vs personal safety) for X/Y evacuated clarity"),
    ("Reproducibility", "GitHub + kaggle_wids_2026_minutes_matter.ipynb + minutes_matter_kaggle.py"),
]

REFERENCES = (
    "References: (1) WiDS Worldwide, WiDS University Datathon 2026, Track 1. "
    "(2) WatchDuty — https://www.watchduty.org/ "
    "(3) NASA FIRMS — https://firms.modaps.eosdis.nasa.gov/ "
    "(4) NIFC / Esri active incident services. "
    "(5) FEMA NRI — https://hazards.fema.gov/nri/ "
    "(6) FEMA NSS Open Shelters. "
    "(7) Open-Meteo — https://open-meteo.com/ "
    "(8) Anthropic Claude API (Flameo). "
    "(9) Supabase, Vercel, Next.js documentation. "
    "Acknowledgments: WiDS Worldwide; WiDS Apprenticeship Pilot; Accenture Song; "
    "UNC Charlotte faculty and advisors; open-data providers; WatchDuty and "
    "first-responder communities."
)

# WiDS font guide — use mid-range sizes
TITLE_PT = 72
AUTHORS_PT = 48
AFFIL_PT = 22
TRACK_PT = 20
SUBTITLE_PT = 18
LOGOS_NOTE_PT = 18
HEADING_PT = 36
BODY_PT = 26
CAPTION_PT = 22
PLACEHOLDER_PT = 19
TABLE_PT = 19
REFS_PT = 18

PURPLE = RGBColor(0x5B, 0x21, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)


def _set_run_font(run, *, name: str = "Calibri", size_pt: float | None = None, bold: bool = False, color=None):
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_heading_para(cell, text: str, size_pt: float = HEADING_PT):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_font(r, size_pt=size_pt, bold=True, color=PURPLE)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(10)


def _add_body_para(cell, text: str, size_pt: float = BODY_PT):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_font(r, size_pt=size_pt, bold=False)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08


def _add_caption_para(cell, text: str):
    p = cell.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size_pt=CAPTION_PT, bold=False)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(4)


def _add_placeholder_para(cell, text: str):
    p = cell.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size_pt=PLACEHOLDER_PT, bold=True, color=GRAY)
    p.paragraph_format.space_after = Pt(6)


def _add_refs_para(cell, text: str):
    p = cell.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size_pt=REFS_PT, bold=False)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05


def _outer_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in list(tblPr):
        if el.tag == qn("w:tblBorders"):
            tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        borders.append(b)
    tblPr.append(borders)
    for el in list(tblPr):
        if el.tag == qn("w:tblLayout"):
            tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for el in list(tblPr):
        if el.tag == qn("w:tblW"):
            tblPr.remove(el)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(int(47 * 1440)))
    tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)


def _set_cell_margins(cell, top=108, bottom=108, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _fill_data_table(table, rows: list[tuple[str, ...]]):
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(str(val))
            _set_run_font(r, size_pt=TABLE_PT, bold=(i == 0))
    _style_inner_table_borders(table)


def _style_inner_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in list(tblPr):
        if el.tag == qn("w:tblBorders"):
            tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "666666")
        borders.append(b)
    tblPr.append(borders)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(48)
    section.page_height = Inches(36)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    outer = doc.add_table(rows=2, cols=3)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    _outer_table_no_borders(outer)

    col_w = Inches(47 / 3)
    for row in outer.rows:
        for cell in row.cells:
            cell.width = col_w
            _set_cell_margins(cell)

    header_cell = outer.rows[0].cells[0]
    header_cell.merge(outer.rows[0].cells[1])
    header_cell.merge(outer.rows[0].cells[2])

    while len(header_cell.paragraphs) > 1:
        el = header_cell.paragraphs[0]._element
        el.getparent().remove(el)

    p = header_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    _set_run_font(r, size_pt=TITLE_PT, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(AUTHORS)
    _set_run_font(r, size_pt=AUTHORS_PT, bold=False)
    p.paragraph_format.space_after = Pt(4)

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(AFFILIATIONS)
    _set_run_font(r, size_pt=AFFIL_PT, bold=False)
    p.paragraph_format.space_after = Pt(4)

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(LOGOS_NOTE)
    _set_run_font(r, size_pt=LOGOS_NOTE_PT, bold=False, color=GRAY)
    p.paragraph_format.space_after = Pt(4)

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TRACK_LINE)
    _set_run_font(r, size_pt=TRACK_PT, bold=True, color=PURPLE)
    p.paragraph_format.space_after = Pt(2)

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    _set_run_font(r, size_pt=SUBTITLE_PT, bold=False)
    p.paragraph_format.space_after = Pt(12)

    c0, c1, c2 = outer.rows[1].cells

    def _drop_extra_empty_paragraphs(cell) -> None:
        while len(cell.paragraphs) > 1 and not (cell.paragraphs[0].text or "").strip():
            el = cell.paragraphs[0]._element
            el.getparent().remove(el)

    _drop_extra_empty_paragraphs(c1)
    _drop_extra_empty_paragraphs(c2)

    # Column 0 — Abstract + Introduction + Figures 1–2 (WiDS: abstract text-only; intro + figs)
    _add_heading_para(c0, "Abstract", HEADING_PT)
    _add_body_para(c0, ABSTRACT, BODY_PT)

    _add_heading_para(c0, "Introduction", HEADING_PT)
    _add_body_para(c0, INTRODUCTION, BODY_PT)
    _add_placeholder_para(c0, FIGURE_PLACEHOLDERS[1])
    _add_caption_para(c0, CAPTIONS[1])
    _add_placeholder_para(c0, FIGURE_PLACEHOLDERS[2])
    _add_caption_para(c0, CAPTIONS[2])

    # Column 1 — Data and Methods + Table 1 + Figures 3–4
    _add_heading_para(c1, "Data and Methods", HEADING_PT)
    _add_body_para(c1, DATA_METHODS, BODY_PT)

    _add_heading_para(c1, "Table 1. Data sources integrated in Minutes Matter", HEADING_PT)
    t1 = c1.add_table(rows=len(TABLE1_ROWS), cols=3)
    _fill_data_table(t1, TABLE1_ROWS)

    _add_placeholder_para(c1, FIGURE_PLACEHOLDERS[3])
    _add_caption_para(c1, CAPTIONS[3])
    _add_placeholder_para(c1, FIGURE_PLACEHOLDERS[4])
    _add_caption_para(c1, CAPTIONS[4])

    # Column 2 — Results + Table 2 + Figures 5–8 + Discussion + References
    _add_heading_para(c2, "Results", HEADING_PT)
    _add_body_para(c2, RESULTS, BODY_PT)

    _add_heading_para(c2, "Table 2. Solution highlights (quantitative + product)", HEADING_PT)
    t2 = c2.add_table(rows=len(TABLE2_ROWS), cols=2)
    _fill_data_table(t2, TABLE2_ROWS)

    for n in (5, 6, 7, 8):
        _add_placeholder_para(c2, FIGURE_PLACEHOLDERS[n])
        _add_caption_para(c2, CAPTIONS[n])

    _add_heading_para(c2, "Discussion", HEADING_PT)
    _add_body_para(c2, DISCUSSION, BODY_PT)

    _add_heading_para(c2, "References and Acknowledgments", HEADING_PT)
    _add_refs_para(c2, REFERENCES)

    return doc


_CHILD_RE = re.compile(r"<w:(top|bottom|left|right|insideH|insideV)\b[^>]*/>")


def fix_tbl_borders_order_in_docx(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    xml = data["word/document.xml"].decode("utf-8")

    def reorder_block(m: re.Match) -> str:
        start_tag = m.group(1)
        inner = m.group(2)
        by_name: dict[str, str] = {}
        for cm in _CHILD_RE.finditer(inner):
            by_name[cm.group(1)] = cm.group(0)
        stripped = _CHILD_RE.sub("", inner).strip()
        order = ("top", "bottom", "left", "right", "insideH", "insideV")
        rebuilt = "".join(by_name[n] for n in order if n in by_name)
        for k, v in by_name.items():
            if k not in order:
                rebuilt += v
        if stripped:
            rebuilt += stripped
        return f"{start_tag}{rebuilt}</w:tblBorders>"

    tb_pat = re.compile(r"(<w:tblBorders\b[^>]*>)(.*?)</w:tblBorders>", re.DOTALL)
    xml2 = tb_pat.sub(reorder_block, xml)
    data["word/document.xml"] = xml2.encode("utf-8")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    path.write_bytes(buf.getvalue())


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    out = root / "docs" / "WiDS_Poster_MinutesMatter.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc = build_document()
    except Exception as e:
        print("python-docx build failed:", e, file=sys.stderr)
        print("Install: pip install python-docx", file=sys.stderr)
        return 1

    doc.save(out)
    fix_tbl_borders_order_in_docx(out)
    print("Wrote", out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
